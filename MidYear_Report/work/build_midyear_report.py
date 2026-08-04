#!/usr/bin/env python3
"""Build the revised five-page mid-year technical report from the partner draft."""

from __future__ import annotations

import csv
import glob
import json
import math
import shutil
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("/Users/andy/Desktop/P4P")
PROJECT = ROOT / "P4P_knowledgeWiki"
PYANSYS = PROJECT / "PyAnsys"
WORK = PROJECT / "MidYear_Report" / "work"
REFERENCE = PROJECT / "MidYear_Report" / "MidYearReport_Partner_Draft_Reference.docx"
OUTPUT = PROJECT / "MidYear_Report" / "MidYearTechnicalReport_LiteratureReviewed.docx"
SWEEP = PYANSYS / "output" / "spiral_enthalpy_sweep_20260725"
FIGURES = WORK / "figures"


def set_run_font(run, name: str = "Times New Roman", size: float | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)


def clear_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 70, start: int = 90, bottom: int = 70, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_twips: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total = sum(widths_twips)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "90")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_twips[index])


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_twips: list[int],
    *,
    font_size: float = 8.3,
    numeric_columns: set[int] | None = None,
) -> None:
    numeric_columns = numeric_columns or set()
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = text
        set_cell_shading(cell, "D9EAF4")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if index in numeric_columns else WD_ALIGN_PARAGRAPH.LEFT
        )
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, size=font_size)
            run.bold = True
    for values in rows:
        cells = table.add_row().cells
        for index, text in enumerate(values):
            cells[index].text = text
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cells[index])
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if index in numeric_columns else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                set_run_font(run, size=font_size)
    set_table_geometry(table, widths_twips)


def add_body(document: Document, text: str, *, bold_lead: str = "") -> None:
    paragraph = document.add_paragraph(style="Body Text")
    paragraph.paragraph_format.keep_together = False
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        lead.bold = True
        set_run_font(lead, size=10.5)
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(rest, size=10.5)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=10.5)


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Caption")
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=9)


def add_figure(
    document: Document,
    path: Path,
    width_inches: float,
    alt_text: str,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    picture = run.add_picture(str(path), width=Inches(width_inches))
    picture._inline.docPr.set("descr", alt_text)
    picture._inline.docPr.set("title", alt_text)


def add_equation(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run, name="Cambria Math", size=10.5)
    run.italic = True


def add_page_break(document: Document) -> None:
    document.add_page_break()


def make_dpm_montage() -> Path:
    FIGURES.mkdir(parents=True, exist_ok=True)
    sources = [
        (
            SWEEP / "plots" / "dpm_particle_tracking" / "spiral_inlet_dpm_particle_tracks_5_micron.png",
            "(a) 5 µm, residence scale 0-18 s",
        ),
        (
            SWEEP / "plots" / "dpm_particle_tracking" / "spiral_inlet_dpm_particle_tracks_168_micron.png",
            "(b) 168 µm, residence scale 0-10.6 s",
        ),
        (
            SWEEP / "plots" / "dpm_particle_tracking" / "spiral_inlet_dpm_particle_tracks_348_micron.png",
            "(c) 348 µm, residence scale 0-4.08 s",
        ),
    ]
    panel_width = 900
    image_height = 675
    label_height = 64
    canvas = Image.new("RGB", (panel_width * 3, image_height + label_height), "white")
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 30)
    except OSError:
        font = ImageFont.load_default()
    draw = ImageDraw.Draw(canvas)
    for index, (path, label) in enumerate(sources):
        image = Image.open(path).convert("RGB")
        image = image.resize((panel_width, image_height), Image.Resampling.LANCZOS)
        x = index * panel_width
        canvas.paste(image, (x, label_height))
        box = draw.textbbox((0, 0), label, font=font)
        text_width = box[2] - box[0]
        draw.text((x + (panel_width - text_width) / 2, 14), label, fill="black", font=font)
        if index:
            draw.line((x, 0, x, image_height + label_height), fill="#B7B7B7", width=3)
    output = FIGURES / "spiral_dpm_particle_tracking_comparison.png"
    canvas.save(output, quality=95)
    return output


def load_case_rows() -> list[dict[str, str]]:
    with (SWEEP / "all_enthalpy_case_summary.csv").open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def load_reference_simulation() -> list[dict[str, str]]:
    path = PYANSYS / "output" / "graph_digitization" / "spiral_inlet_reference_digitized_points.csv"
    with path.open(newline="", encoding="utf-8-sig") as handle:
        return [row for row in csv.DictReader(handle) if row["series"] == "Simulation"]


def load_plot_rows() -> list[dict[str, str]]:
    path = SWEEP / "plots" / "spiral_inlet_output_steam_quality_plot_data.csv"
    with path.open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def nearest_reference_quality(velocity: float, rows: list[dict[str, str]]) -> float:
    nearest = min(rows, key=lambda row: abs(float(row["x_mps"]) - velocity))
    return float(nearest["steam_quality_pct"])


def final_residual_rows() -> list[list[str]]:
    rows: list[list[str]] = []
    for path_text in sorted(
        glob.glob(str(SWEEP / "case_*_1500iter_residual_history.csv"))
    ):
        path = Path(path_text)
        with path.open(newline="", encoding="utf-8-sig") as handle:
            values = list(csv.DictReader(handle))[-1]
        case_number = path.name.split("_")[1]
        rows.append(
            [
                case_number,
                f"{float(values['continuity']):.3f}",
                f"{float(values['vf-phase-2']):.2e}",
                f"{max(float(values[k]) for k in ('x-velocity', 'y-velocity', 'z-velocity')):.2e}",
            ]
        )
    return rows


def set_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")

    body = styles["Body Text"]
    body.font.size = Pt(10.5)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.line_spacing = 1.08
    body.paragraph_format.space_after = Pt(4)

    for name, size, after in (
        ("Heading 1", 14, 5),
        ("Heading 2", 11.5, 3),
        ("Heading 3", 10.5, 2),
    ):
        style = styles[name]
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(5)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if "MidYear Title" not in styles:
        title = styles.add_style("MidYear Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        title = styles["MidYear Title"]
    title.font.name = "Times New Roman"
    title.font.size = Pt(18)
    title.font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)


def set_header_footer(document: Document) -> None:
    for section in document.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(0.92)
        section.different_first_page_header_footer = False
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(16.0))
        left = paragraph.add_run("Mid-Year Technical Report")
        set_run_font(left, size=9)
        paragraph.add_run("\t")
        right = paragraph.add_run("Geothermal Separator CFD")
        set_run_font(right, size=9)


def add_title(document: Document) -> None:
    paragraph = document.add_paragraph(style="MidYear Title")
    run = paragraph.add_run(
        "CFD Analysis of a Geothermal Steam-Water Separator:\n"
        "Spiral-Inlet Reconstruction, DPM Tracking and Wall-Film Sensitivity"
    )
    set_run_font(run, size=18)
    run.bold = True
    scope = document.add_paragraph(style="Body Text")
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scope.paragraph_format.space_after = Pt(5)
    run = scope.add_run(
        "Technical results from controlled ANSYS Fluent simulations and automated PyFluent post-processing"
    )
    set_run_font(run, size=9.5)
    run.italic = True


def build() -> None:
    if not REFERENCE.exists():
        raise FileNotFoundError(REFERENCE)
    if not SWEEP.exists():
        raise FileNotFoundError(SWEEP)

    dpm_montage = make_dpm_montage()
    quality_figure = SWEEP / "plots" / "spiral_inlet_output_steam_quality.png"
    wall_film_figure = SWEEP / "plots" / "wall_film" / "spiral_inlet_wall_film_thickness.png"

    document = Document(REFERENCE)
    clear_body(document)
    set_styles(document)
    set_header_footer(document)

    case_rows = load_case_rows()
    plot_rows = load_plot_rows()
    reference_rows = load_reference_simulation()

    add_title(document)
    document.add_paragraph("Numerical model and automated workflow", style="Heading 1")
    add_body(
        document,
        "The study used staged CFD comparisons to quantify liquid carryover in a vertical "
        "bottom-outlet cyclone separator [2]. The first model family reconstructed six operating "
        "conditions from Purnanto et al. (2013) on the available rectangular 90° spiral-inlet "
        "geometry [1]. Because the paper does not publish CAD-ready scroll geometry or its exact "
        "nine-injection table, this is a controlled reconstruction rather than an exact duplicate. "
        "A second family introduced split steam-liquid inlet and Eulerian wall-film (EWF) "
        "treatments to isolate model-form sensitivity."
    )
    add_caption(document, "Table 1. Principal numerical settings and controlled model differences.")
    add_table(
        document,
        ["Item", "Spiral-inlet reconstruction", "Split-inlet / EWF family"],
        [
            ["Mesh", "2,964,593 tetrahedral cells", "7,601,261 cells"],
            ["Solver", "Steady, pressure based; Mixture model", "Steady, pressure based; Mixture + DPM/EWF"],
            ["Turbulence", "RNG k-ε; standard wall functions", "RNG k-ε; standard wall functions"],
            ["Thermal scope", "Energy disabled; no flashing", "Energy disabled; no flashing"],
            ["Inlet", "Single two-phase mass-flow inlet", "Separate steam and liquid mass-flow regions"],
            ["Outlet / lower boundary", "1.12 MPa pressure outlet / DPM trap", "Pressure outlet / closed lower wall"],
            ["Discrete phase", "Nine Harwell-derived bins; one-way DPM", "Deterministic/stochastic DPM and EWF branches"],
        ],
        [1900, 3600, 3600],
        font_size=7.9,
    )
    document.add_paragraph("Operating-condition and injection reconstruction", style="Heading 2")
    add_body(
        document,
        "Each enthalpy label selected the steam-liquid mass split reported by Purnanto; enthalpy "
        "was not solved as a thermal boundary condition [1]. The Harwell correlation set each "
        "case's Sauter/median diameter scale, while nine normalized diameter and mass bins were "
        "reconstructed from the paper's Figure 5. The spiral baseline's native surface-injection "
        "convention was preserved: inert liquid-water particles were released from the inlet "
        "using the local face-normal direction."
    )
    add_equation(
        document,
        "xₛₐ = 1.91 Dₜ (Re⁰·¹ / We⁰·⁶)(ρg / ρl)⁰·⁶,     xₘₑd = 1.42 xₛₐ"
    )
    add_body(
        document,
        "The automated sequence loaded a fresh baseline, set and read back both phase mass "
        "flows, updated all nine injections, hybrid-initialised the carrier flow, completed "
        "1500 verified iterations in 25-iteration blocks, saved a pre-DPM state, performed DPM "
        "tracking, and exported injection-level fates. Early residual stopping was disabled. "
        "Each case required nine populated injection rows and fate closure within 0.2%; the "
        "completed six-case sweep therefore required 54 rows."
    )
    add_equation(
        document,
        "ṁseparated = ṁtrapped + ṁincomplete,     quality_Purnanto = ṁsteam /(ṁsteam + ṁescaped) × 100%"
    )

    add_page_break(document)
    document.add_paragraph("Completed spiral-inlet enthalpy sweep", style="Heading 1")
    add_body(
        document,
        "All six operating conditions completed 1500 carrier-flow iterations on 15 Fluent ranks, "
        "followed by per-injection DPM tracking. Escaped liquid remained between 0.01655 and "
        "0.02667 kg/s. Following Purnanto's treatment, all incomplete trajectories were assigned "
        "to the separated/trapped category and only escaped liquid was counted as carryover. This "
        "gives 99.96681% to 99.97949% steam quality. These values apply the published fate "
        "interpretation but are not independently audited outlet qualities. The present points "
        "form a nearly flat band and do not reproduce the 26.85 m/s dip in Purnanto's digitised "
        "spiral-inlet simulation series or its high-velocity calculation decline [1]."
    )
    add_figure(
        document,
        quality_figure,
        6.25,
        "Scatter plot comparing calculated outlet steam quality for six spiral-inlet CFD cases "
        "with digitised Purnanto calculation, simulation, and correlation results.",
    )
    add_caption(
        document,
        "Figure 1. Present six-case spiral-inlet CFD results compared with digitised Purnanto "
        "Figure 20 calculation/simulation points and the Lazalde-Crabtree correlation. Black "
        "points treat trapped and incomplete trajectories as separated, following Purnanto; the "
        "quality axis is inverted to match the source."
    )
    result_rows: list[list[str]] = []
    for summary, plotted in zip(case_rows, plot_rows):
        velocity = float(plotted["inlet_velocity_ms"])
        project_quality = float(plotted["steam_quality_pct"])
        reference_quality = nearest_reference_quality(velocity, reference_rows)
        result_rows.append(
            [
                summary["case"].replace("Case ", ""),
                summary["enthalpy_kJkg"],
                f"{velocity:.2f}",
                f"{float(summary['escaped_kgs']):.5f}",
                f"{float(summary['incomplete_kgs']):.3f}",
                f"{project_quality:.4f}",
                f"{project_quality-reference_quality:+.4f}",
            ]
        )
    add_caption(document, "Table 2. Completed spiral-inlet DPM outcomes and comparison with the nearest digitised Figure 20 simulation point.")
    add_table(
        document,
        ["Case", "Condition", "v (m/s)", "Escaped\n(kg/s)", "Incomplete\n(kg/s)", "Quality\n(%)", "Δ vs ref.\n(pp)"],
        result_rows,
        [650, 1450, 1050, 1300, 1400, 1300, 1200],
        font_size=7.7,
        numeric_columns={0, 2, 3, 4, 5, 6},
    )
    add_body(
        document,
        "Case 4 is the largest trend mismatch: the project predicts 99.9795% at 26.87 m/s, "
        "whereas the nearest digitised Purnanto simulation point is approximately 99.8297%. "
        "Agreement in magnitude for the other five points does not establish validation because "
        "the reference trend is not recovered, the numerator is prescribed inlet steam flow, and "
        "incomplete trajectories are assigned to the separated category rather than resolved to "
        "a physical endpoint."
    )

    add_page_break(document)
    document.add_paragraph("Size-resolved DPM behaviour", style="Heading 1")
    add_body(
        document,
        "Across all six spiral cases, completed outlet escape occurred only in the finest "
        "injection. The 5 µm-class tracks remain closely coupled to the rotating carrier flow, "
        "populate the upper recirculation region and exhibit the longest plotted residence-time "
        "scale. Increasing diameter shortens the displayed residence-time range and shifts the "
        "dominant fate toward bottom trapping, consistent with increasing inertial separation."
    )
    add_figure(
        document,
        dpm_montage,
        6.45,
        "Three particle-track panels showing residence time for nominal 5, 168, and 348 micrometre "
        "droplets in the spiral-inlet separator.",
    )
    add_caption(
        document,
        "Figure 2. Particle-residence-time tracks for nominal 5, 168 and 348 µm injections. "
        "Colour scales differ between panels and therefore indicate within-panel residence time, not a shared quantitative scale."
    )
    add_caption(document, "Table 3. Selected Case 4 (1600 kJ/kg) injection-level fate accounting.")
    case4_path = SWEEP / "case_4_1600_injection_results.csv"
    with case4_path.open(newline="", encoding="utf-8-sig") as handle:
        case4_all = {row["injection_name"]: row for row in csv.DictReader(handle)}
    selected = [
        ("injection-5-micron", "5"),
        ("injection-168-micron", "168"),
        ("injection-348-micron", "348"),
    ]
    selected_rows = []
    for key, label in selected:
        row = case4_all[key]
        injected = float(row["injected_mass_flow_kgs"])
        selected_rows.append(
            [
                label,
                f"{float(row['diameter_mm'])*1000:.1f}",
                f"{injected:.4f}",
                f"{float(row['escaped_kgs']):.5f}",
                f"{float(row['trapped_kgs']):.4f}",
                f"{float(row['incomplete_kgs']):.4f}",
                f"{100*float(row['incomplete_kgs'])/injected:.1f}",
            ]
        )
    add_table(
        document,
        ["Nominal class\n(µm)", "Actual d\n(µm)", "Injected\n(kg/s)", "Escaped\n(kg/s)", "Trapped\n(kg/s)", "Incomplete\n(kg/s)", "Incomplete\n(%)"],
        selected_rows,
        [1150, 1050, 1300, 1300, 1300, 1400, 1400],
        font_size=8.0,
        numeric_columns=set(range(7)),
    )
    add_body(
        document,
        "The Case 4 finest bin is an important numerical exception: 0.01655 kg/s escaped, but "
        "0.1362 kg/s remained incomplete, so incomplete mass was over eight times the escaped "
        "mass for that bin. No completed escape was recorded for the 168 or 348 µm classes. "
        "The 348 µm class instead delivered 22.43 kg/s to the bottom trap, while 0.9562 kg/s "
        "remained incomplete. The raw CSV retains incomplete as a separate Fluent fate for "
        "traceability; for the reported Purnanto-consistent quality, it is added to trapped mass "
        "and treated as separated."
    )

    add_page_break(document)
    document.add_paragraph("Eulerian wall-film mechanism screening", style="Heading 1")
    add_body(
        document,
        "The wall-film family partitioned the liquid input between Eulerian inlet liquid and DPM "
        "parcels before enabling deposition and drainage, avoiding double counting. Fluent's EWF "
        "formulation supports DPM collection, splash, edge separation and shear-driven stripping "
        "[3]. These branches are project mechanism screens, not Purnanto replication cases."
    )
    add_figure(
        document,
        wall_film_figure,
        4.55,
        "ANSYS Fluent wall-film thickness contour on the vertical separator and spiral inlet, "
        "with the largest film accumulation at the inlet bend.",
    )
    add_caption(
        document,
        "Figure 3. Wall-film thickness on the separator and spiral inlet. The plotted scale spans "
        "0-3.94×10⁻⁴ m and shows the largest local accumulation at the inlet bend and first vessel-wall impact region."
    )
    add_caption(document, "Table 4. Captured film quantities from the mechanism-screening branches.")
    add_table(
        document,
        ["EWF branch", "Inventory\n(kg)", "Maximum\n(mm)", "Area average\n(µm)"],
        [
            ["Clean deposition", "0.07150", "0.152", "1.211"],
            ["Splash only", "0.07431", "0.164", "1.259"],
            ["Edge separation only", "0.05639", "0.123", "0.955"],
            ["Stripping only", "0.05440", "0.121", "0.921"],
            ["Combined mechanisms", "0.05668", "0.125", "0.960"],
        ],
        [3600, 1800, 1800, 1900],
        font_size=8.1,
        numeric_columns={1, 2, 3},
    )
    add_body(
        document,
        "Relative to the clean deposition control, splash increased inventory by 3.9% and "
        "maximum thickness by 7.9%, whereas edge separation and stripping reduced inventory by "
        "21.1% and 23.9%, respectively. In the combined EWF branch, enabling global DPM "
        "interaction increased inventory, maximum thickness and area-average thickness by 41%, "
        "and increased derived mean film speed from 0.0533 to 0.0684 m/s (+28%). The interaction-on "
        "state also increased coarse-class splash events from 20 to 44, edge-separation events "
        "from 120 to 188, and introduced seven stripping events. Values are transcribed from the "
        "project's captured Fluent snapshots and were not rerun within the six-case sweep."
    )
    add_body(
        document,
        "Event counts are not equivalent to re-entrained liquid mass. They demonstrate that the "
        "selected wall model changes the assigned liquid pathway, but a quantitative carryover "
        "claim requires the generated secondary parcels to be mass-accounted and compared from "
        "matched initial film states and physical times."
    )

    add_page_break(document)
    document.add_paragraph("Technical interpretation and limitations", style="Heading 1")
    document.add_paragraph("Iteration completion versus convergence", style="Heading 2")
    add_body(
        document,
        "Every spiral case contains 1500 residual-history points and passed the workflow’s setup, "
        "save and DPM mass-closure checks. However, fixed iteration count is not convergence "
        "evidence. Final scaled continuity residuals remain between 0.145 and 0.229, far above "
        "the configured 10⁻⁴ criterion, while momentum residuals are approximately 1.1×10⁻⁵ to "
        "1.7×10⁻⁵. The carrier fields are iteration-complete but not strictly converged."
    )
    add_caption(document, "Table 5. Final scaled residuals after 1500 iterations.")
    add_table(
        document,
        ["Case", "Continuity", "Liquid volume fraction", "Maximum momentum"],
        final_residual_rows(),
        [1200, 2100, 2900, 2900],
        font_size=8.4,
        numeric_columns={0, 1, 2, 3},
    )
    document.add_paragraph("Treatment of incomplete trajectories", style="Heading 2")
    incomplete_fractions = [
        100 * float(row["incomplete_kgs"]) / float(row["liquid_mass_flow_kgs"])
        for row in case_rows
    ]
    add_body(
        document,
        f"Incomplete liquid increases from {min(incomplete_fractions):.1f}% to "
        f"{max(incomplete_fractions):.1f}% across the six spiral cases and is substantially "
        "larger than escaped liquid. Fluent defines an incomplete fate numerically: the maximum "
        "tracking-step limit was reached [4]. Purnanto encountered the same issue and assumed "
        "incomplete particles were separated to obtain a quality estimate [1]. The present report "
        "uses the same convention: all incomplete mass is combined with trapped mass, and only "
        "escaped mass is counted as liquid carryover. The separate incomplete field is retained "
        "to quantify tracking resolution and sensitivity to the maximum-step setting."
    )
    document.add_paragraph("Evidence boundary", style="Heading 2")
    add_body(
        document,
        "The completed sweep establishes a reproducible numerical pipeline and a size-selective "
        "carryover mechanism, but it does not validate separator efficiency. The exact nine-bin "
        "mass allocation is reconstructed because Purnanto did not publish the original injection "
        "table or CAD-ready scroll definition; the model is isothermal and omits flashing, breakup "
        "and coalescence; steam quality uses prescribed inlet steam and Purnanto's incomplete-as-"
        "separated convention rather than an audited outlet vapour monitor; and the reference "
        "trend is not recovered. The defensible mid-year result is therefore sensitivity to "
        "droplet size and wall treatment, not validated efficiency."
    )
    document.add_paragraph("Technical verification required before final validation", style="Heading 2")
    add_body(
        document,
        "The next numerical checks are to extend the carrier solutions while monitoring integral "
        "outlet mass flow and pressure drop, export incomplete-particle endpoints, repeat DPM "
        "tracking across step-length and stochastic-dispersion settings, and audit quality against "
        "Fluent outlet vapour flow. One representative case should also test mesh and RSM "
        "sensitivity; experiment-backed cyclone work uses RSM-DPM for anisotropic rotating flow, "
        "although its air-water values are not transferable here [5]. EWF interaction-off/on "
        "cases must be rerun from a common film state and physical time."
    )
    source = document.add_paragraph(style="Caption")
    source.paragraph_format.space_before = Pt(5)
    run = source.add_run(
        "References: [1] Purnanto, Zarrouk & Cater (2013), IPENZ Transactions 40, 1-10. "
        "[2] Zarrouk & Purnanto (2014), Geothermics 53, 236-254, doi:10.1016/j.geothermics.2014.05.009. "
        "[3] Ansys Fluent Theory Guide 2025 R2, §§17.1-17.2.1."
    )
    set_run_font(run, size=7.5)
    source2 = document.add_paragraph(style="Caption")
    source2.paragraph_format.space_before = Pt(0)
    source2.paragraph_format.space_after = Pt(0)
    run = source2.add_run(
        "[4] Ansys Fluent User's Guide 2025 R2, §§24.2.3 and 25.9.3. "
        "[5] Chen et al. (2025), Processes 13, 3732, doi:10.3390/pr13113732."
    )
    set_run_font(run, size=7.5)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
